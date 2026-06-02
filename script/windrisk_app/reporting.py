"""Report generation utilities."""

import io

import numpy as np
import pandas as pd
from docx import Document


def generate_docx_report(df_results_raw, total_assets, total_portfolio_value_nok):
    """Generate a DOCX report from batch analysis results."""
    document = Document()

    df_calc = df_results_raw.copy()
    df_calc["Asset Value (Numeric)"] = df_calc["Asset Value"].apply(
        lambda x: float(str(x).replace(",", "")) if str(x) != "N/A" else np.nan
    )
    df_calc["Average Annual Loss (Numeric)"] = df_calc["Average Annual Loss (NOK)"].apply(
        lambda x: float(str(x).replace(",", "")) if str(x) != "N/A" else np.nan
    )
    df_calc["Risk Score (Numeric)"] = df_calc["Risk Index (No Adapt)"].apply(
        lambda x: float(str(x)) if str(x) != "N/A" else np.nan
    )

    num_assets = total_assets
    total_exposure = df_calc["Asset Value (Numeric)"].sum()
    average_aal = df_calc["Average Annual Loss (Numeric)"].mean()
    total_expected_annual_loss = df_calc["Average Annual Loss (Numeric)"].sum()

    high_risk_threshold = 80
    high_risk_assets_count = df_calc[df_calc["Risk Score (Numeric)"] > high_risk_threshold].shape[0]
    percent_high_risk = (high_risk_assets_count / num_assets * 100) if num_assets > 0 else 0

    document.add_heading("Portfolio Wind Risk Assessment Report", level=1)
    document.add_paragraph(f"Date: {pd.Timestamp.now().strftime('%Y-%m-%d')}")

    document.add_heading("1. Executive Summary", level=2)
    document.add_paragraph(
        "This report presents an assessment of wind-related risks for the aggregated property portfolio, "
        "as provided via the uploaded Excel file. Leveraging detailed wind hazard data for Bergen, "
        "building vulnerability models, and the Klawa Risk Index, this analysis quantifies potential financial losses "
        "and categorizes assets by risk level under various adaptation scenarios."
    )

    document.add_heading("2. Portfolio Overview", level=2)
    document.add_paragraph(f"Total Number of Assets Analyzed: {num_assets}")
    document.add_paragraph(f"Total Exposure (Sum of Asset Values): {total_exposure:,.2f} NOK")
    document.add_paragraph(f"Average Annual Loss (AAL) per Asset: {average_aal:,.2f} NOK")
    document.add_paragraph(f"Total Expected Annual Loss (Portfolio): {total_expected_annual_loss:,.2f} NOK")
    document.add_paragraph(
        f"Percentage of Assets in High-Risk Category (Score > {high_risk_threshold}): {percent_high_risk:.2f}%"
    )

    document.add_heading("3. Batch Analysis Results", level=2)
    document.add_paragraph(
        "This table provides a detailed breakdown of the wind risk assessment for each individual property in the uploaded portfolio."
    )

    df_results_for_report = df_results_raw[
        [
            "Address",
            "Asset Value",
            "Risk Index (No Adapt)",
            "Average Annual Loss (NOK)",
            "Map Risk Zone (No Adapt)",
        ]
    ].rename(
        columns={
            "Risk Index (No Adapt)": "Risk Score",
            "Map Risk Zone (No Adapt)": "Adaptation Level",
        }
    )

    table = document.add_table(rows=1, cols=len(df_results_for_report.columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for idx, col_name in enumerate(df_results_for_report.columns):
        hdr_cells[idx].text = str(col_name)

    for _, row in df_results_for_report.iterrows():
        row_cells = table.add_row().cells
        for idx, col_name in enumerate(df_results_for_report.columns):
            value = row[col_name]
            if isinstance(value, (bytes, bytearray)):
                try:
                    value = value.decode("utf-8")
                except UnicodeDecodeError:
                    value = value.decode("latin-1", errors="replace")
            else:
                value = str(value)

            if col_name in ["Asset Value", "Risk Score", "AAL (NOK)"]:
                try:
                    if "NOK" in col_name and value != "N/A":
                        cleaned_value = value.replace(",", "")
                        value = f"{float(cleaned_value):,.2f}"
                    elif col_name == "Risk Score" and value != "N/A":
                        value = f"{float(value):.2f}"
                    elif value != "N/A":
                        value = f"{float(value):.2f}"
                except ValueError:
                    pass

            row_cells[idx].text = value

    document.add_heading("4. Scenario Notes", level=2)
    document.add_paragraph(
        "Scores calculated using baseline wind scenario with Klawa impact function and normalized to max pooled risk (100)."
    )

    document.add_heading("5. Risk Methodology Summary", level=2)
    document.add_paragraph(
        "Wind risk is calculated by combining local wind hazard data (return period wind speeds) with building-specific "
        "vulnerability curves, which model the damage ratio as a function of wind speed."
    )

    document.add_heading("6. Contact Information & Disclaimer", level=2)
    document.add_paragraph(
        "This is a prototype assessment. For advisory or regulatory use, please contact Climate Intelligence for calibrated estimates."
    )

    doc_buffer = io.BytesIO()
    document.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer.getvalue()
