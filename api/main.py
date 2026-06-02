from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.gzip import GZipMiddleware
from fastapi.responses import Response
import xarray as xr
import numpy as np
import pandas as pd
import matplotlib.colors as mcolors
from scipy.spatial import cKDTree
import osmnx as ox
from shapely.geometry import box
from geopy.geocoders import Nominatim
from geopy.exc import GeocoderTimedOut, GeocoderServiceError
from docx import Document
import io
import logging
import os
import pickle
import asyncio
from functools import partial
from pathlib import Path

app = FastAPI(title="WindRisk API")

app.add_middleware(GZipMiddleware, minimum_size=1000)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Paths are resolved relative to this file so uvicorn can be run from any directory
_HERE = Path(__file__).parent
wind_data_path    = str(_HERE / '../data/bergen_return_period_winds_1985_2020_debugged.nc')
klawa_base_data_path = str(_HERE / '../data/klawa_risk_index.nc')
klawa_v99_data_path  = str(_HERE / '../data/klawa_risk_index_v99.nc')
klawa_v995_data_path = str(_HERE / '../data/klawa_risk_index_v995.nc')
OSM_CACHE_PATH    = str(_HERE / '../cache/buildings_osm.pkl')

VULNERABILITY_PARAMETERS = {
    'Weak':     {'k': 0.4894, 'v0': 29.9949},
    'Moderate': {'k': 0.2887, 'v0': 41.3318},
    'Strong':   {'k': 0.1821, 'v0': 53.9558}
}

RED_ZONE_THRESHOLD_PERC    = 90
ORANGE_ZONE_THRESHOLD_PERC = 70
YELLOW_ZONE_THRESHOLD_PERC = 30
GEOMETRY_SIMPLIFICATION_TOLERANCE = 0.00001

MAP_LEVEL_TO_WIND_VAR = {
    'No Adaptation':     'wind_98th_percentile_speed',
    'Medium Adaptation': 'wind_99th_percentile_speed',
    'High Adaptation':   'wind_995th_percentile_speed'
}

klawa_data_sources_paths = {
    'no_adapt':     {'path': klawa_base_data_path, 'wind_var_name': MAP_LEVEL_TO_WIND_VAR['No Adaptation']},
    'medium_adapt': {'path': klawa_v99_data_path,  'wind_var_name': MAP_LEVEL_TO_WIND_VAR['Medium Adaptation']},
    'high_adapt':   {'path': klawa_v995_data_path, 'wind_var_name': MAP_LEVEL_TO_WIND_VAR['High Adaptation']}
}

# --- Globals ---
pydeck_buildings_gdf   = None
full_buildings_gdf     = None
map_data_json_bytes    = None
map_data_light_bytes   = None   # centroid-only payload (~10x smaller)
center_lat = 0
center_lon = 0
display_return_periods = []
building_centroids     = None
all_return_periods     = []
buildings_tree         = None


def get_risk_color_rgba(klawa_index_norm):
    if pd.isna(klawa_index_norm):
        return [128, 128, 128, 0]
    if klawa_index_norm >= RED_ZONE_THRESHOLD_PERC:
        hex_color = '#DE2D26'
    elif klawa_index_norm >= ORANGE_ZONE_THRESHOLD_PERC:
        hex_color = '#FB6A4A'
    elif klawa_index_norm >= YELLOW_ZONE_THRESHOLD_PERC:
        hex_color = '#FFD700'
    else:
        hex_color = '#D4EDDA'
    return [int(c * 255) for c in mcolors.to_rgba(hex_color)]


def geocode_address(address, user_agent='bergen_wind_app'):
    geolocator = Nominatim(user_agent=user_agent)
    try:
        return geolocator.geocode(address, timeout=10)
    except (GeocoderTimedOut, GeocoderServiceError) as e:
        logger.error(f"Geocoding error: {e}")
        return None


def calculate_aal(building_data, resilience_type, all_available_return_periods, asset_value_nok):
    if asset_value_nok is None or pd.isna(asset_value_nok):
        return np.nan
    rp_damage_pairs = []
    for rp in all_available_return_periods:
        damage_ratio = building_data.get(f'damage_ratio_{resilience_type.lower()}_rp{rp}')
        if pd.notna(damage_ratio):
            rp_damage_pairs.append((rp, damage_ratio))
    if not rp_damage_pairs:
        return np.nan
    rp_damage_pairs.sort(key=lambda x: x[0])
    sorted_rps    = [p[0] for p in rp_damage_pairs]
    sorted_damage = [p[1] for p in rp_damage_pairs]
    probabilities = [1.0] + [1.0 / rp for rp in sorted_rps] + [0.0]
    damages       = [sorted_damage[0]] + sorted_damage + [0.0]
    aal_value = 0.0
    for i in range(len(probabilities) - 1):
        p1, p2 = probabilities[i], probabilities[i + 1]
        d1, d2 = damages[i], damages[i + 1]
        if p1 >= p2:
            aal_value += 0.5 * (d1 + d2) * (p1 - p2)
    return aal_value * asset_value_nok


def generate_docx_report(df_results_raw, total_assets, total_portfolio_value_nok):
    document = Document()
    df_calc = pd.DataFrame(df_results_raw)
    df_calc['Asset Value (Numeric)'] = df_calc['Asset Value'].apply(
        lambda x: float(str(x).replace(',', '')) if str(x) != 'N/A' else np.nan)
    df_calc['Average Annual Loss (Numeric)'] = df_calc['Average Annual Loss (NOK)'].apply(
        lambda x: float(str(x).replace(',', '')) if str(x) != 'N/A' else np.nan)
    df_calc['Risk Score (Numeric)'] = df_calc['Risk Index (No Adapt)'].apply(
        lambda x: float(str(x)) if str(x) != 'N/A' else np.nan)

    total_exposure             = df_calc['Asset Value (Numeric)'].sum()
    average_aal                = df_calc['Average Annual Loss (Numeric)'].mean()
    total_expected_annual_loss = df_calc['Average Annual Loss (Numeric)'].sum()
    high_risk_threshold        = 80
    high_risk_assets_count     = df_calc[df_calc['Risk Score (Numeric)'] > high_risk_threshold].shape[0]
    percent_high_risk          = (high_risk_assets_count / total_assets * 100) if total_assets > 0 else 0

    document.add_heading('Portfolio Wind Risk Assessment Report', level=1)
    document.add_paragraph(f'Date: {pd.Timestamp.now().strftime("%Y-%m-%d")}')
    document.add_heading('1. Executive Summary', level=2)
    document.add_paragraph("This report presents an assessment of wind-related risks for the aggregated property portfolio...")
    document.add_heading('2. Portfolio Overview', level=2)
    document.add_paragraph(f"Total Number of Assets Analyzed: {total_assets}")
    document.add_paragraph(f"Total Exposure (Sum of Asset Values): {total_exposure:,.2f} NOK")
    document.add_paragraph(f"Average Annual Loss (AAL) per Asset: {average_aal:,.2f} NOK")
    document.add_paragraph(f"Total Expected Annual Loss (Portfolio): {total_expected_annual_loss:,.2f} NOK")
    document.add_paragraph(f"Percentage of Assets in High-Risk Category (Score > {high_risk_threshold}): {percent_high_risk:.2f}%")

    document.add_heading('3. Batch Analysis Results', level=2)
    cols_to_show = ['Address', 'Asset Value', 'Risk Index (No Adapt)', 'Average Annual Loss (NOK)', 'Map Risk Zone (No Adapt)']
    df_report = df_calc[cols_to_show].rename(columns={
        'Risk Index (No Adapt)': 'Risk Score',
        'Map Risk Zone (No Adapt)': 'Adaptation Level'
    })
    table = document.add_table(rows=1, cols=len(df_report.columns))
    table.style = 'Table Grid'
    for i, col_name in enumerate(df_report.columns):
        table.rows[0].cells[i].text = str(col_name)
    for _, row in df_report.iterrows():
        row_cells = table.add_row().cells
        for i, col_name in enumerate(df_report.columns):
            value = row[col_name]
            value = value.decode('utf-8', errors='replace') if isinstance(value, (bytes, bytearray)) else str(value)
            try:
                if 'NOK' in col_name and value not in ('N/A', 'nan'):
                    value = f"{float(value.replace(',', '')):,.2f}"
                elif col_name == 'Risk Score' and value not in ('N/A', 'nan'):
                    value = f"{float(value):.2f}"
                elif value not in ('N/A', 'nan'):
                    value = f"{float(value):.2f}"
            except ValueError:
                pass
            row_cells[i].text = value

    document.add_heading('4. Scenario Notes', level=2)
    document.add_paragraph("Scores calculated using baseline wind scenario with Klawa impact function and normalized...")
    document.add_heading('5. Risk Methodology Summary', level=2)
    document.add_paragraph("Wind risk is calculated by combining local wind hazard data...")
    document.add_heading('6. Contact Information & Disclaimer', level=2)
    document.add_paragraph("This is a prototype assessment...")

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf.getvalue()


@app.on_event("startup")
def load_data():
    global pydeck_buildings_gdf, full_buildings_gdf, map_data_json_bytes, map_data_light_bytes, center_lat, center_lon
    global display_return_periods, building_centroids, all_return_periods, buildings_tree

    logger.info("Loading wind and klawa datasets...")
    wind_data = xr.open_dataset(wind_data_path)
    wind_lons = wind_data.x.values
    wind_lats = wind_data.y.values
    min_lon, max_lon = wind_lons.min(), wind_lons.max()
    min_lat, max_lat = wind_lats.min(), wind_lats.max()

    klawa_and_wind_grids_raw_xarray = {}
    klawa_lons, klawa_lats = None, None
    for key, info in klawa_data_sources_paths.items():
        ds = xr.open_dataset(info['path'])
        klawa_and_wind_grids_raw_xarray[key] = {
            'klawa_risk_index':      ds['klawa_risk_index'].squeeze(),
            'wind_percentile_speed': ds[info['wind_var_name']].squeeze()
        }
        if klawa_lons is None:
            klawa_lons = ds.x.values
            klawa_lats = ds.y.values

    bbox_polygon = box(min_lon, min_lat, max_lon, max_lat)
    tags = {"building": ["residential", "house", "apartments", "flats", "detached",
                         "semidetached_house", "terrace", "bungalow", "farm", "shed", "cabin"]}

    # --- OSM cache: skip the Overpass API call on warm restarts ---
    if os.path.exists(OSM_CACHE_PATH):
        logger.info("Loading OSM buildings from disk cache...")
        buildings = pickle.load(open(OSM_CACHE_PATH, "rb"))
    else:
        logger.info("Fetching OSM buildings from Overpass API (first run only)...")
        buildings = ox.features_from_polygon(bbox_polygon, tags)
        buildings = buildings[buildings.geometry.type == 'Polygon'].copy()
        buildings = buildings.to_crs(epsg=4326)
        buildings['geometry'] = buildings.geometry.simplify(
            GEOMETRY_SIMPLIFICATION_TOLERANCE, preserve_topology=True)
        buildings = buildings[~buildings.geometry.is_empty].copy()
        centroids_tmp = buildings.geometry.centroid
        valid_mask = centroids_tmp.apply(
            lambda p: p is not None and np.isfinite(p.x) and np.isfinite(p.y))
        buildings = buildings[valid_mask][['geometry']].copy()
        os.makedirs(os.path.dirname(OSM_CACHE_PATH), exist_ok=True)
        pickle.dump(buildings, open(OSM_CACHE_PATH, "wb"))
        logger.info(f"OSM buildings cached to {OSM_CACHE_PATH}")

    building_centroids = buildings.geometry.centroid
    valid_mask = building_centroids.apply(
        lambda p: p is not None and np.isfinite(p.x) and np.isfinite(p.y))
    buildings = buildings[valid_mask].copy()
    building_centroids = building_centroids[valid_mask]
    building_coords = np.array([[p.x, p.y] for p in building_centroids])

    buildings_tree = cKDTree(building_coords)

    all_return_periods_val  = wind_data.return_period.values
    all_return_periods      = [int(x) for x in all_return_periods_val]
    display_return_periods  = [int(x) for x in all_return_periods_val if x != 1]

    # Assign wind speeds per building via KDTree lookup
    wind_lon_mesh, wind_lat_mesh = np.meshgrid(wind_lons, wind_lats)
    wind_tree = cKDTree(np.column_stack([wind_lon_mesh.flatten(), wind_lat_mesh.flatten()]))
    _, wind_indices = wind_tree.query(building_coords)
    wind_rows, wind_cols = np.unravel_index(wind_indices, (len(wind_lats), len(wind_lons)))

    for rp in all_return_periods_val:
        ws_grid = wind_data['ReturnPeriodWindSpeed'].sel(return_period=rp).squeeze()
        buildings[f'wind_speed_rp{rp}'] = ws_grid.isel(
            y=xr.DataArray(wind_rows, dims="i"),
            x=xr.DataArray(wind_cols, dims="i")
        ).values

    # Vectorized vulnerability curves (sigmoid) — replaces per-row Python .apply()
    for b_type, params in VULNERABILITY_PARAMETERS.items():
        k, v0 = params['k'], params['v0']
        for rp in all_return_periods_val:
            ws = buildings[f'wind_speed_rp{rp}'].values
            buildings[f'damage_ratio_{b_type.lower()}_rp{rp}'] = np.where(
                np.isfinite(ws),
                1.0 / (1.0 + np.exp(-k * (ws - v0))),
                0.0
            )

    # Assign klawa values
    klawa_lon_mesh, klawa_lat_mesh = np.meshgrid(klawa_lons, klawa_lats)
    klawa_tree = cKDTree(np.column_stack([klawa_lon_mesh.flatten(), klawa_lat_mesh.flatten()]))
    _, klawa_indices = klawa_tree.query(building_coords)
    klawa_rows, klawa_cols = np.unravel_index(klawa_indices, (len(klawa_lats), len(klawa_lons)))

    all_raw_klawa_values = []
    for key in klawa_data_sources_paths:
        klawa_da  = klawa_and_wind_grids_raw_xarray[key]['klawa_risk_index']
        wind_p_da = klawa_and_wind_grids_raw_xarray[key]['wind_percentile_speed']
        buildings[f'klawa_risk_index_raw_{key}'] = klawa_da.isel(
            y=xr.DataArray(klawa_rows, dims="i"),
            x=xr.DataArray(klawa_cols, dims="i")
        ).values
        buildings[f'wind_speed_percentile_{key}'] = wind_p_da.isel(
            y=xr.DataArray(klawa_rows, dims="i"),
            x=xr.DataArray(klawa_cols, dims="i")
        ).values
        all_raw_klawa_values.extend(buildings[f'klawa_risk_index_raw_{key}'].dropna().tolist())

    global_min_klawa = np.nanmin(all_raw_klawa_values) if all_raw_klawa_values else 0.0
    global_max_klawa = np.nanmax(all_raw_klawa_values) if all_raw_klawa_values else 1.0
    if global_max_klawa == global_min_klawa:
        global_max_klawa += 1e-6

    for key in klawa_data_sources_paths:
        raw_col  = f'klawa_risk_index_raw_{key}'
        norm_col = f'normalized_klawa_risk_index_{key}'
        zone_col = f'risk_zone_{key}'
        color_col = f'risk_color_rgba_{key}'

        valid = buildings[raw_col].dropna()
        normed = np.clip(((valid - global_min_klawa) / (global_max_klawa - global_min_klawa)) * 100, 0, 100)
        buildings[norm_col] = np.nan
        buildings.loc[valid.index, norm_col] = normed

        buildings[zone_col] = 'Green'
        vi = buildings[norm_col].dropna().index
        buildings.loc[vi[buildings.loc[vi, norm_col].between(YELLOW_ZONE_THRESHOLD_PERC, ORANGE_ZONE_THRESHOLD_PERC, inclusive='left')], zone_col] = 'Yellow'
        buildings.loc[vi[buildings.loc[vi, norm_col].between(ORANGE_ZONE_THRESHOLD_PERC, RED_ZONE_THRESHOLD_PERC, inclusive='left')], zone_col] = 'Orange'
        buildings.loc[vi[buildings.loc[vi, norm_col] >= RED_ZONE_THRESHOLD_PERC], zone_col] = 'Red'
        buildings[color_col] = buildings[norm_col].apply(get_risk_color_rgba)

    minx, miny, maxx, maxy = buildings.total_bounds
    center_lat = (miny + maxy) / 2
    center_lon = (minx + maxx) / 2

    pydeck_buildings_gdf = buildings[[
        'geometry',
        'risk_color_rgba_no_adapt',
        'risk_color_rgba_medium_adapt',
        'risk_color_rgba_high_adapt'
    ]].copy()

    buildings['id_obj'] = [str(idx) for idx in buildings.index]
    full_buildings_gdf = buildings

    # Pre-serialize full polygon payload
    map_data_json_bytes = pydeck_buildings_gdf.to_json().encode()

    # Build lightweight centroid payload — lon/lat + 3 color arrays, no geometry
    import json as _json
    centroids = buildings.geometry.centroid
    light_features = [
        {
            "lon": round(c.x, 6),
            "lat": round(c.y, 6),
            "c0": buildings.at[idx, 'risk_color_rgba_no_adapt'],
            "c1": buildings.at[idx, 'risk_color_rgba_medium_adapt'],
            "c2": buildings.at[idx, 'risk_color_rgba_high_adapt'],
        }
        for idx, c in centroids.items()
    ]
    map_data_light_bytes = _json.dumps(light_features).encode()
    logger.info(
        f"Data loaded. Full: {len(map_data_json_bytes)//1024} KB  "
        f"Light: {len(map_data_light_bytes)//1024} KB"
    )


@app.get("/api/status")
def status():
    return {"status": "ok", "ready": full_buildings_gdf is not None}


@app.get("/api/map_data")
def get_map_data():
    if map_data_json_bytes is None:
        raise HTTPException(status_code=503, detail="Data not loaded yet")
    return Response(
        content=map_data_json_bytes,
        media_type="application/json",
        headers={"Cache-Control": "public, max-age=3600"}
    )


@app.get("/api/map_data_light")
def get_map_data_light():
    if map_data_light_bytes is None:
        raise HTTPException(status_code=503, detail="Data not loaded yet")
    return Response(
        content=map_data_light_bytes,
        media_type="application/json",
        headers={"Cache-Control": "public, max-age=3600"}
    )


@app.get("/api/analyze")
async def analyze_address(address: str, asset_value: float = None):
    if full_buildings_gdf is None:
        raise HTTPException(status_code=503, detail="Data not loaded yet")

    loop = asyncio.get_event_loop()
    location = await loop.run_in_executor(None, partial(geocode_address, address))
    if not location:
        raise HTTPException(status_code=404, detail="Address not found")

    _, indices = buildings_tree.query([[location.longitude, location.latitude]])
    building = full_buildings_gdf.iloc[indices[0]]

    result = {
        "address": location.address,
        "building_id": str(building.name),
        "lat": location.latitude,
        "lon": location.longitude,
        "klawa": {
            k: (float(building[f'normalized_klawa_risk_index_{v}'])
                if pd.notna(building[f'normalized_klawa_risk_index_{v}']) else None)
            for k, v in [("No Adaptation", "no_adapt"),
                         ("Medium Adaptation", "medium_adapt"),
                         ("High Adaptation", "high_adapt")]
        },
        "wind_speed": {
            k: (float(building[f'wind_speed_percentile_{v}'])
                if pd.notna(building[f'wind_speed_percentile_{v}']) else None)
            for k, v in [("No Adaptation", "no_adapt"),
                         ("Medium Adaptation", "medium_adapt"),
                         ("High Adaptation", "high_adapt")]
        },
        "risk_zone": {
            "No Adaptation":     building['risk_zone_no_adapt'],
            "Medium Adaptation": building['risk_zone_medium_adapt'],
            "High Adaptation":   building['risk_zone_high_adapt']
        },
        "aal": {},
        "scenarios": []
    }

    asset_val_total = asset_value * 1_000_000 if asset_value else 0
    for res in ["Weak", "Moderate", "Strong"]:
        result["aal"][res] = calculate_aal(building, res, all_return_periods, asset_val_total) if asset_val_total else None

    for rp in display_return_periods:
        robj = {"rp": rp, "wind_speed": float(building[f'wind_speed_rp{rp}'])}
        for res in ["Weak", "Moderate", "Strong"]:
            dr = building[f'damage_ratio_{res.lower()}_rp{rp}']
            dr_val = float(dr) if pd.notna(dr) else None
            robj[f'dr_{res.lower()}'] = dr_val
            robj[f'damage_{res.lower()}'] = (dr_val * asset_val_total) if (dr_val and asset_val_total) else None
        result["scenarios"].append(robj)

    return result


@app.post("/api/batch")
async def batch_analysis(file: UploadFile = File(...)):
    if full_buildings_gdf is None:
        raise HTTPException(status_code=503, detail="Data not loaded yet")

    df_uploaded = pd.read_excel(file.file)
    required_cols = ['Address', 'Monetary Asset Value (million NOK)']
    if not all(col in df_uploaded.columns for col in required_cols):
        raise HTTPException(status_code=400, detail=f"Excel must contain: {required_cols}")

    batch_results = []
    total_portfolio_asset_value = 0.0
    loop = asyncio.get_event_loop()
    geolocator = Nominatim(user_agent='bergen_wind_app_batch')

    for _, row in df_uploaded.iterrows():
        address = row['Address']
        asset_value_millions = float(row['Monetary Asset Value (million NOK)'])
        total_portfolio_asset_value += asset_value_millions * 1_000_000

        try:
            location = await loop.run_in_executor(
                None, partial(geolocator.geocode, address, timeout=10))
        except Exception:
            location = None

        if location:
            _, indices = buildings_tree.query([[location.longitude, location.latitude]])
            building = full_buildings_gdf.iloc[indices[0]]
            dr_rp100 = building.get('damage_ratio_moderate_rp100', np.nan)
            calculated_damage = (dr_rp100 * asset_value_millions * 1_000_000
                                 if pd.notna(dr_rp100) else np.nan)
            aal_nok = calculate_aal(building, 'Moderate', all_return_periods,
                                    asset_value_millions * 1_000_000)
            batch_results.append({
                'Address': address,
                'Asset Value': f"{asset_value_millions * 1_000_000:,.2f}",
                'Risk Index (No Adapt)': building['normalized_klawa_risk_index_no_adapt'],
                'Risk Index (Medium Adapt)': building['normalized_klawa_risk_index_medium_adapt'],
                'Risk Index (High Adapt)': building['normalized_klawa_risk_index_high_adapt'],
                'Damage Ratio (RP 100 yr, Moderate)': dr_rp100.item() if pd.notna(dr_rp100) else None,
                'Calculated Damage (NOK)': calculated_damage,
                'Average Annual Loss (NOK)': aal_nok,
                'Map Risk Zone (No Adapt)': building['risk_zone_no_adapt']
            })
        else:
            batch_results.append({
                'Address': address, 'Asset Value': 'N/A',
                'Risk Index (No Adapt)': 'N/A', 'Risk Index (Medium Adapt)': 'N/A',
                'Risk Index (High Adapt)': 'N/A', 'Average Annual Loss (NOK)': 'N/A',
                'Map Risk Zone (No Adapt)': 'N/A', 'Error': 'Could not geocode'
            })

    return {"results": batch_results, "total_value": total_portfolio_asset_value}


@app.post("/api/report")
async def get_report(data: dict):
    try:
        docx_bytes = generate_docx_report(
            data.get('results', []), len(data.get('results', [])), data.get('total_value', 0))
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Portfolio_Wind_Risk_Report.docx"}
        )
    except Exception as e:
        logger.error(str(e))
        raise HTTPException(status_code=500, detail="Error generating report")
