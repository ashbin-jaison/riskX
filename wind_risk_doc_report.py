import streamlit as st
import xarray as xr
import numpy as np
import pandas as pd
import matplotlib.colors as mcolors
import matplotlib.pyplot as plt
from scipy.spatial import cKDTree
import geopandas as gpd
import osmnx as ox
from shapely.geometry import box
import pydeck as pdk
from geopy.geocoders import Nominatim
from geopy.exc import GeocoderTimedOut, GeocoderServiceError
# Removed fpdf import
from docx import Document # Import python-docx library
from docx.shared import Inches # For setting image size if needed
import io # For handling byte streams for DOCX download

# --- Configuration Parameters ---
# IMPORTANT: Adjust these paths to your actual NetCDF file locations
wind_data_path = 'bergen_return_period_winds_1985_2020_debugged.nc'
klawa_base_data_path = 'klawa_risk_index.nc'
klawa_v99_data_path = 'klawa_risk_index_v99.nc'
klawa_v995_data_path = 'klawa_risk_index_v995.nc'

# Logo file paths - ADJUST THESE TO YOUR EXACT FILE LOCATIONS
LOGO_PATH_1 = 'cl_logo_tp.png'

# Vulnerability curve parameters for different building types: D(U10) = 1 / (1 + exp(-k * (U10 - v0)))
# L is implicitly 1.0 (100% damage ratio) for all these functions.
VULNERABILITY_PARAMETERS = {
    'Weak': {'k': 0.4894, 'v0': 29.9949},
    'Moderate': {'k': 0.2887, 'v0': 41.3318},
    'Strong': {'k': 0.1821, 'v0': 53.9558}
}

# Risk zone percentile thresholds applied to normalized Klawa Risk Index (0-100)
RED_ZONE_THRESHOLD_PERC = 90    # Risk Index >= 90
ORANGE_ZONE_THRESHOLD_PERC = 70 # Risk Index >= 70 and < 90
YELLOW_ZONE_THRESHOLD_PERC = 30 # Risk Index >= 30 and < 70
# Green: Risk Index < 30

# The return period selected for displaying damage ratio in tooltips
SELECTED_RP_FOR_DAMAGE_PLOT = 100

# Geometry simplification tolerance (in degrees). Adjust as needed.
GEOMETRY_SIMPLIFICATION_TOLERANCE = 0.00001

# Mapping for adaptation level display name to internal column suffix
MAP_LEVEL_TO_SUFFIX = {
    'No Adaptation': 'no_adapt',
    'Medium Adaptation': 'medium_adapt', # Simplified name
    'High Adaptation': 'high_adapt'      # Simplified name
}
# Mapping for adaptation level display name to corresponding wind percentile variable name in NC files
MAP_LEVEL_TO_WIND_VAR = {
    'No Adaptation': 'wind_98th_percentile_speed',
    'Medium Adaptation': 'wind_99th_percentile_speed', # Simplified name
    'High Adaptation': 'wind_995th_percentile_speed'    # Simplified name
}


# --- Helper Functions ---

# Initialize colormap globally (not directly used for map coloring anymore)
CMAP_DAMAGE = plt.cm.get_cmap('YlOrRd')

@st.cache_data
def vulnerability_curve(wind_speed, k_val, v0_val):
    """Sigmoid vulnerability curve function: D(U10) = 1 / (1 + exp(-k * (U10 - v0)))"""
    if not np.isfinite(wind_speed):
        return 0.0
    return 1.0 / (1 + np.exp(-k_val * (wind_speed - v0_val)))

def get_risk_color_rgba(klawa_index_norm):
    """Converts normalized Klawa index to an RGBA color array based on defined thresholds."""
    if klawa_index_norm >= RED_ZONE_THRESHOLD_PERC:
        hex_color = '#DE2D26' # Red
    elif klawa_index_norm >= ORANGE_ZONE_THRESHOLD_PERC:
        hex_color = '#FB6A4A' # Orange
    elif klawa_index_norm >= YELLOW_ZONE_THRESHOLD_PERC:
        hex_color = '#FFD700' # Yellow
    else:
        hex_color = '#D4EDDA' # Green
    rgba_color = mcolors.to_rgba(hex_color)
    return [int(c * 255) for c in rgba_color]

# Helper for odometer/gauge color based on normalized Klawa Index
def get_gauge_color_hex(klawa_index_norm):
    if klawa_index_norm >= RED_ZONE_THRESHOLD_PERC:
        return '#DE2D26' # Red
    elif klawa_index_norm >= ORANGE_ZONE_THRESHOLD_PERC:
        return '#FB6A4A' # Orange
    elif klawa_index_norm >= YELLOW_ZONE_THRESHOLD_PERC:
        return '#FFD700' # Yellow
    else:
        return '#D4EDDA' # Green

# --- Geocoding Function ---
@st.cache_data
def geocode_address(address, user_agent='bergen_wind_app'):
    """Geocodes an address using Nominatim."""
    geolocator = Nominatim(user_agent=user_agent)
    try:
        location = geolocator.geocode(address, timeout=10)
        return location
    except (GeocoderTimedOut, GeocoderServiceError) as e:
        st.error(f"Geocoding error: {e}. Please try again or refine the address.")
        return None

# Function to calculate Average Annual Loss (AAL)
def calculate_aal(building_data, resilience_type, all_available_return_periods, asset_value_nok):
    """
    Calculates Average Annual Loss (AAL) for a building given its damage ratios
    across various return periods. Uses a trapezoidal approximation of the
    Exceedance Probability vs. Damage curve.
    """
    if asset_value_nok is None or pd.isna(asset_value_nok):
        return np.nan # Return NaN if no asset value provided

    # Get damage ratios for all return periods for the given resilience type
    rp_damage_pairs = []
    for rp in all_available_return_periods:
        col_name = f'damage_ratio_{resilience_type.lower()}_rp{rp}'
        damage_ratio = building_data.get(col_name) # Use .get to handle potential missing columns gracefully
        if pd.notna(damage_ratio):
            rp_damage_pairs.append((rp, damage_ratio))

    if not rp_damage_pairs:
        return np.nan # No valid damage data to calculate AAL

    # Sort by return period ascending
    rp_damage_pairs.sort(key=lambda x: x[0])

    # Extract sorted return periods and damage ratios
    sorted_rps = [item[0] for item in rp_damage_pairs]
    sorted_damage_ratios = [item[1] for item in rp_damage_pairs]

    # Calculate exceedance probabilities P = 1/RP
    # Add a conceptual point (P=1, D=damage_at_lowest_RP) and (P=0, D=0) for proper integration
    # This is a common practice for AAL curves starting from 100% probability (RP=1) down to 0% probability (RP=infinity)
    probabilities = [1.0] + [1.0 / rp for rp in sorted_rps] + [0.0]
    damages = [sorted_damage_ratios[0] if sorted_damage_ratios else 0.0] + sorted_damage_ratios + [0.0] # Assume damage at RP=1 is first damage, and damage at RP=inf is 0

    aal_value = 0.0
    for i in range(len(probabilities) - 1):
        p1 = probabilities[i]
        p2 = probabilities[i+1]
        d1 = damages[i]
        d2 = damages[i+1]
        
        # Area of trapezoid: 0.5 * (sum of heights) * base
        # Here, heights are damages, base is change in probability
        # Since probabilities are descending, p1 - p2 is positive
        if p1 >= p2: # Ensure probability is decreasing
             aal_value += 0.5 * (d1 + d2) * (p1 - p2)
        else: # Handle cases where sorting might not capture all nuances or if data has non-monotonic probabilities
            # This should ideally not happen if probabilities are derived simply from 1/RP and RPs are sorted.
            # For robustness, we could consider absolute difference or log-linear interpolation.
            # For now, let's just add a warning or handle as 0 for this segment if P increases.
            st.warning(f"Non-monotonic probability detected in AAL calculation: {p1} -> {p2}. Segment skipped.")


    # Multiply by asset value to get AAL in monetary terms
    return aal_value * asset_value_nok


# --- Main Data Processing Pipeline (Cached) ---

@st.cache_data
def load_and_process_data(wind_data_path, klawa_data_sources_paths, vuln_params, simplify_tolerance):
    """
    Loads wind data, all Klawa risk index data files, OSM buildings,
    performs spatial join, calculates damage ratios, and assigns risk zones.
    """
    try:
        wind_data = xr.open_dataset(wind_data_path)
    except FileNotFoundError:
        st.error(f"Error: Wind data file not found at {wind_data_path}. Please check the path.")
        st.stop()
    except Exception as e:
        st.error(f"Error loading wind data: {e}")
        st.stop()

    wind_lons = wind_data.x.values
    wind_lats = wind_data.y.values

    min_lon, max_lon = wind_lons.min(), wind_lons.max()
    min_lat, max_lat = wind_lats.min(), wind_lats.max()

    north, south, east, west = max_lat, min_lat, max_lon, min_lon

    # Load Klawa Risk Index Data and associated Wind Percentiles for all adaptation levels
    klawa_and_wind_grids_raw_xarray = {} # Stores raw xarray DataArrays for both klawa and wind percentiles
    klawa_lons = None
    klawa_lats = None

    for key, info in klawa_data_sources_paths.items():
        try:
            temp_klawa_ds = xr.open_dataset(info['path'])
            klawa_and_wind_grids_raw_xarray[key] = {
                'klawa_risk_index': temp_klawa_ds['klawa_risk_index'].squeeze(), # Variable name is always 'klawa_risk_index'
                'wind_percentile_speed': temp_klawa_ds[info['wind_var_name']].squeeze()
            }
            if klawa_lons is None: # Use coordinates from the first Klawa file loaded
                klawa_lons = temp_klawa_ds.x.values
                klawa_lats = temp_klawa_ds.y.values
        except FileNotFoundError:
            st.error(f"Error: Klawa risk index or wind percentile file not found at {info['path']}. Please check the path.")
            st.stop()
        except KeyError as ke:
            st.error(f"Error: Variable '{ke}' not found in {info['path']}. Please confirm variable names.")
            st.stop()
        except Exception as e:
            st.error(f"Error loading Klawa/Wind data for '{key}': {e}")
            st.stop()

    # Load OSM Buildings
    bbox_polygon = box(min_lon, min_lat, max_lon, max_lat)
    tags = {"building": ["residential", "house", "apartments", "flats",
                          "detached", "semidetached_house", "terrace",
                          "bungalow", "farm", "shed", "cabin"]}
    try:
        buildings = ox.features_from_polygon(bbox_polygon, tags)
        buildings = buildings[buildings.geometry.type == 'Polygon'].copy()
        buildings = buildings.to_crs(epsg=4326)
    except Exception as e:
        st.error(f"Error loading OSM data: {e}. Please ensure OSMnx is updated and working correctly.")
        st.stop()

    if buildings.empty:
        st.warning("No residential buildings found. Check OSM tags or the derived bounding box.")
        return None, None, None, None, None, None, None # Added None for all_return_periods

    buildings['geometry'] = buildings.geometry.simplify(simplify_tolerance, preserve_topology=True)
    
    # Filter out buildings that resulted in empty geometries after simplification
    buildings = buildings[~buildings.geometry.is_empty].copy()

    if buildings.empty:
        st.warning("All residential buildings were filtered out or simplified to empty geometries. Cannot proceed with analysis.")
        return None, None, None, None, None, None, None

    building_centroids = buildings.geometry.centroid
    
    # Filter out any non-finite (NaN, Inf) centroids and update buildings DataFrame
    # This ensures building_coords contains only valid numbers for KDTree query
    valid_centroids_mask = building_centroids.apply(lambda p: p is not None and p.x is not None and p.y is not None and np.isfinite(p.x) and np.isfinite(p.y))
    buildings = buildings[valid_centroids_mask].copy() # Update buildings to only include those with valid centroids
    building_centroids = building_centroids[valid_centroids_mask] # Update building_centroids series

    if buildings.empty: # Re-check after filtering invalid centroids
        st.warning("No valid building centroids found after filtering. Cannot proceed with spatial analysis.")
        return None, None, None, None, None, None, None

    building_coords = np.array([[p.x, p.y] for p in building_centroids])

    all_return_periods = wind_data.return_period.values
    # Filter out 1-year return period from the list for display in table (original 'return_periods' variable)
    display_return_periods = [rp for rp in all_return_periods if rp != 1]

    # --- Define wind_tree and wind_grid_rows/cols ONCE here ---
    # Prepare KDTrees for spatial joins
    wind_lon_mesh, wind_lat_mesh = np.meshgrid(wind_lons, wind_lats)
    wind_points = np.column_stack([wind_lon_mesh.flatten(), wind_lat_mesh.flatten()])
    wind_tree = cKDTree(wind_points) # wind_tree is defined here

    distances, indices = wind_tree.query(building_coords)
    wind_grid_shape = (len(wind_lats), len(wind_lons))
    wind_grid_rows, wind_grid_cols = np.unravel_index(indices, wind_grid_shape)
    # --- End of definition for wind_tree and wind_grid_rows/cols ---

    # Process all return periods for data loading and damage calculations
    for rp in all_return_periods:
        buildings[f'wind_speed_rp{rp}'] = np.nan
        current_wind_speeds_data = wind_data['ReturnPeriodWindSpeed'].sel(return_period=rp).squeeze()
        
        # Assign values using .loc and the index of the buildings GeoDataFrame
        buildings.loc[buildings.index, f'wind_speed_rp{rp}'] = current_wind_speeds_data.isel(
            y=xr.DataArray(wind_grid_rows, dims="temp_dim"),
            x=xr.DataArray(wind_grid_cols, dims="temp_dim")
        ).values


    # Calculate Damage Ratios for all vulnerability types (for individual analysis)
    for b_type, params in vuln_params.items():
        k_val = params['k']
        v0_val = params['v0']
        for rp in all_return_periods:
            buildings[f'damage_ratio_{b_type.lower()}_rp{rp}'] = buildings[f'wind_speed_rp{rp}'].apply(
                lambda x: vulnerability_curve(x, k_val, v0_val)
            )

    # --- Define klawa_tree and klawa_grid_rows/cols ONCE here ---
    klawa_lon_mesh, klawa_lat_mesh = np.meshgrid(klawa_lons, klawa_lats)
    klawa_points = np.column_stack([klawa_lon_mesh.flatten(), klawa_lat_mesh.flatten()])
    klawa_tree = cKDTree(klawa_points) # klawa_tree is defined here

    distances_klawa, indices_klawa = klawa_tree.query(building_coords)
    klawa_grid_shape = (len(klawa_lats), len(klawa_lons))
    klawa_grid_rows, klawa_grid_cols = np.unravel_index(indices_klawa, klawa_grid_shape)
    # --- End of definition for klawa_tree and klawa_grid_rows/cols ---

    all_raw_klawa_values_for_global_min_max = []

    for key in klawa_data_sources_paths.keys():
        klawa_grid_da = klawa_and_wind_grids_raw_xarray[key]['klawa_risk_index']
        wind_percentile_da = klawa_and_wind_grids_raw_xarray[key]['wind_percentile_speed']

        raw_klawa_col_name = f'klawa_risk_index_raw_{key}'
        # This is where the wind percentile speed column name is defined in the DataFrame
        raw_wind_percentile_col_name = f'wind_speed_percentile_{key}'

        # Populate raw klawa and wind speed columns
        buildings.loc[buildings.index, raw_klawa_col_name] = klawa_grid_da.isel(
            y=xr.DataArray(klawa_grid_rows, dims="temp_dim"),
            x=xr.DataArray(klawa_grid_cols, dims="temp_dim")
        ).values
        buildings.loc[buildings.index, raw_wind_percentile_col_name] = wind_percentile_da.isel(
            y=xr.DataArray(klawa_grid_rows, dims="temp_dim"),
            x=xr.DataArray(klawa_grid_cols, dims="temp_dim")
        ).values
        
        # Collect only non-NaN klawa values for global min/max calculation
        all_raw_klawa_values_for_global_min_max.extend(buildings[raw_klawa_col_name].dropna().tolist())

    # --- Calculate GLOBAL min/max for Klawa Risk Index across ALL raw values ---
    if not all_raw_klawa_values_for_global_min_max:
        global_min_klawa = 0.0
        global_max_klawa = 1.0
        st.warning("No valid Klawa risk index values found to calculate global min/max. Using default range (0-1).")
    else:
        global_min_klawa = np.nanmin(all_raw_klawa_values_for_global_min_max)
        global_max_klawa = np.nanmax(all_raw_klawa_values_for_global_min_max)

    if global_max_klawa == global_min_klawa:
        global_max_klawa += 1e-6 # Add a tiny epsilon to prevent division by zero

    # --- Normalize all Klawa Risk Indices using the GLOBAL min/max ---
    for key in klawa_data_sources_paths.keys():
        raw_klawa_col_name = f'klawa_risk_index_raw_{key}'
        norm_klawa_col_name = f'normalized_klawa_risk_index_{key}'
        risk_zone_col_name = f'risk_zone_{key}'
        risk_color_col_name = f'risk_color_rgba_{key}'

        valid_raw_klawa_values = buildings[raw_klawa_col_name].dropna()
        normalized_klawa_values = ((valid_raw_klawa_values - global_min_klawa) / (global_max_klawa - global_min_klawa)) * 100
        normalized_klawa_values = np.clip(normalized_klawa_values, 0, 100)
        
        buildings[norm_klawa_col_name] = np.nan
        buildings.loc[valid_raw_klawa_values.index, norm_klawa_col_name] = normalized_klawa_values

        buildings[risk_zone_col_name] = 'Green'
        valid_normalized_indices = buildings[norm_klawa_col_name].dropna().index

        buildings.loc[(buildings.index.isin(valid_normalized_indices)) &
                      (buildings[norm_klawa_col_name] >= YELLOW_ZONE_THRESHOLD_PERC) &
                      (buildings[norm_klawa_col_name] < ORANGE_ZONE_THRESHOLD_PERC), risk_zone_col_name] = 'Yellow'
        buildings.loc[(buildings.index.isin(valid_normalized_indices)) &
                      (buildings[norm_klawa_col_name] >= ORANGE_ZONE_THRESHOLD_PERC) &
                      (buildings[norm_klawa_col_name] < RED_ZONE_THRESHOLD_PERC), risk_zone_col_name] = 'Orange'
        buildings.loc[(buildings.index.isin(valid_normalized_indices)) &
                      (buildings[norm_klawa_col_name] >= RED_ZONE_THRESHOLD_PERC), risk_zone_col_name] = 'Red'

        buildings[risk_color_col_name] = buildings[norm_klawa_col_name].apply(
            lambda x: get_risk_color_rgba(x) if pd.notna(x) else [128, 128, 128, 0] # Gray and transparent for NaN
        )

    if not buildings.empty:
        minx, miny, maxx, maxy = buildings.total_bounds
        center_lat = (miny + maxy) / 2
        center_lon = (minx + maxx) / 2
    else:
        center_lat = (north + south) / 2
        center_lon = (east + west) / 2

    # Prepare tooltip string dynamically using the values from the current row
    buildings['tooltip'] = buildings.apply(
        lambda row: f"<b>Building ID:</b> {row.name}<br>"
                    f"<b>Damage Ratio (RP {SELECTED_RP_FOR_DAMAGE_PLOT} yr, Moderate):</b> {(row[f'damage_ratio_moderate_rp{SELECTED_RP_FOR_DAMAGE_PLOT}'] * 100):.2f}%<br>"
                    f"<b>Risk Index (No Adapt):</b> {row['normalized_klawa_risk_index_no_adapt']:.2f}<br>" # Changed to Risk Index
                    f"<b>Risk Index (Medium Adapt):</b> {row['normalized_klawa_risk_index_medium_adapt']:.2f}<br>" # Changed to Risk Index
                    f"<b>Risk Index (High Adapt):</b> {row['normalized_klawa_risk_index_high_adapt']:.2f}<br>" # Changed to Risk Index
                    f"<b>Map Risk Zone (No Adapt):</b> {row['risk_zone_no_adapt']}<br>"
                    f"<b>Map Risk Zone (Medium Adapt):</b> {row['risk_zone_medium_adapt']}<br>"
                    f"<b>Map Risk Zone (High Adapt):</b> {row['risk_zone_high_adapt']}<br>"
                    f"<b>98th Percentile Wind Speed:</b> {row['wind_speed_percentile_no_adapt']:.2f} m/s<br>"
                    f"<b>99th Percentile Wind Speed:</b> {row['wind_speed_percentile_medium_adapt']:.2f} m/s<br>"
                    f"<b>99.5th Percentile Wind Speed:</b> {row['wind_speed_percentile_high_adapt']:.2f} m/s",
        axis=1
    )

    # Ensure all necessary columns for Pydeck are included
    pydeck_buildings_gdf = buildings[[
        'geometry',
        'risk_color_rgba_no_adapt',
        'risk_color_rgba_medium_adapt',
        'risk_color_rgba_high_adapt',
        'tooltip'
    ]].copy()

    return pydeck_buildings_gdf, buildings, center_lat, center_lon, display_return_periods, building_centroids, all_return_periods

# --- DOCX Report Generation Function ---
def generate_docx_report(df_results_raw, total_assets, total_portfolio_value_nok):
    document = Document()

    # --- Portfolio Summary Calculations ---
    # Create a temporary DataFrame for numerical calculations, handling 'N/A'
    df_calc = df_results_raw.copy()
    
    # Convert relevant columns to numeric, coercing errors to NaN
    df_calc['Asset Value (Numeric)'] = df_calc['Asset Value'].apply(lambda x: float(str(x).replace(',', '')) if str(x) != 'N/A' else np.nan)
    df_calc['Average Annual Loss (Numeric)'] = df_calc['Average Annual Loss (NOK)'].apply(lambda x: float(str(x).replace(',', '')) if str(x) != 'N/A' else np.nan)
    df_calc['Risk Score (Numeric)'] = df_calc['Risk Index (No Adapt)'].apply(lambda x: float(str(x)) if str(x) != 'N/A' else np.nan)


    num_assets = total_assets # Already passed
    total_exposure = df_calc['Asset Value (Numeric)'].sum()
    average_aal = df_calc['Average Annual Loss (Numeric)'].mean()
    total_expected_annual_loss = df_calc['Average Annual Loss (Numeric)'].sum()

    high_risk_threshold = 80 # Score > 80
    high_risk_assets_count = df_calc[df_calc['Risk Score (Numeric)'] > high_risk_threshold].shape[0]
    percent_high_risk = (high_risk_assets_count / num_assets * 100) if num_assets > 0 else 0

    # --- Document Content ---
    document.add_heading('Portfolio Wind Risk Assessment Report', level=1)
    document.add_paragraph(f'Date: {pd.Timestamp.now().strftime("%Y-%m-%d")}')
    
    document.add_heading('1. Executive Summary', level=2)
    document.add_paragraph(
        "This report presents an assessment of wind-related risks for the aggregated property portfolio, "
        "as provided via the uploaded Excel file. Leveraging detailed wind hazard data for Bergen, "
        "building vulnerability models, and the Klawa Risk Index, this analysis quantifies potential financial losses "
        "and categorizes assets by risk level under various adaptation scenarios. The primary objective is to "
        "provide a comprehensive understanding of the portfolio's exposure to wind risk, identify high-risk assets, "
        "and highlight the benefits of different adaptation strategies, enabling informed decision-making for risk "
        "mitigation and strategic planning."
    )

    document.add_heading('2. Portfolio Overview', level=2)
    document.add_paragraph(f"Total Number of Assets Analyzed: {num_assets}")
    document.add_paragraph(f"Total Exposure (Sum of Asset Values): {total_exposure:,.2f} NOK")
    document.add_paragraph(f"Average Annual Loss (AAL) per Asset: {average_aal:,.2f} NOK")
    document.add_paragraph(f"Total Expected Annual Loss (Portfolio): {total_expected_annual_loss:,.2f} NOK")
    document.add_paragraph(f"Percentage of Assets in High-Risk Category (Score > {high_risk_threshold}): {percent_high_risk:.2f}%")

    document.add_heading('3. Batch Analysis Results', level=2)
    document.add_paragraph(
        "This table provides a detailed breakdown of the wind risk assessment for each individual property in the uploaded portfolio."
    )

    # Prepare DataFrame for report table columns
    df_results_for_report = df_results_raw[[
        'Address',
        'Asset Value',
        'Risk Index (No Adapt)', # Will be renamed to 'Risk Score'
        'Average Annual Loss (NOK)',
        'Map Risk Zone (No Adapt)' # Will be renamed to 'Adaptation Level'
    ]].rename(columns={
        'Risk Index (No Adapt)': 'Risk Score',
        'Map Risk Zone (No Adapt)': 'Adaptation Level'
    })


    # Add table to the document
    table = document.add_table(rows=1, cols=len(df_results_for_report.columns))
    table.style = 'Table Grid' # Apply a basic table style

    # Add table header
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df_results_for_report.columns):
        hdr_cells[i].text = str(col_name) # Ensure header is string

    # Add data rows
    for index, row in df_results_for_report.iterrows():
        row_cells = table.add_row().cells
        for i, col_name in enumerate(df_results_for_report.columns):
            value = row[col_name]
            # Ensure value is string for docx
            if isinstance(value, (bytes, bytearray)):
                try:
                    value = value.decode('utf-8')
                except UnicodeDecodeError:
                    value = value.decode('latin-1', errors='replace')
            else:
                value = str(value)

            # Special formatting for numerical columns
            if col_name in ['Asset Value', 'Risk Score', 'AAL (NOK)']: # Updated col names for formatting
                try:
                    if 'NOK' in col_name and value != 'N/A':
                        # Clean value by removing any existing commas before converting to float
                        cleaned_value = value.replace(',', '')
                        value = f"{float(cleaned_value):,.2f}"
                    elif col_name == 'Risk Score' and value != 'N/A':
                        value = f"{float(value):.2f}"
                    elif value != 'N/A': # Catch other general numerical columns
                         value = f"{float(value):.2f}"
                except ValueError:
                    pass # Keep as N/A if conversion fails

            row_cells[i].text = value
            
    # --- Scenario Notes ---
    document.add_heading('4. Scenario Notes', level=2)
    document.add_paragraph(
        "Scores calculated using baseline wind scenario with Klawa impact function and normalized to max pooled risk (100)."
    )

    # --- Risk Methodology Summary ---
    document.add_heading('5. Risk Methodology Summary', level=2)
    document.add_paragraph(
        "Wind risk is calculated by combining local wind hazard data (return period wind speeds) with building-specific "
        "vulnerability curves, which model the damage ratio as a function of wind speed. The Klawa Risk Index, "
        "an empirically derived measure, further refines the risk assessment by integrating additional factors beyond "
        "just wind speed, offering a more holistic view of localized wind exposure. The resulting risk scores are "
        "normalized to a scale of 0-100 based on the maximum observed risk within the analyzed portfolio, facilitating "
        "clear categorization into risk zones (Green, Yellow, Orange, Red)."
    )

    # --- Contact Info / Disclaimer ---
    document.add_heading('6. Contact Information & Disclaimer', level=2)
    document.add_paragraph(
        "This is a prototype assessment. For advisory or regulatory use, please contact Climate Intelligence for calibrated estimates."
    )


    # Save document to a bytes buffer
    doc_buffer = io.BytesIO()
    document.save(doc_buffer)
    doc_buffer.seek(0) # Rewind the buffer
    return doc_buffer.getvalue()


# --- Streamlit Application ---

st.set_page_config(
    page_title="Bergen Wind Risk Map",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Define paths for the cached data loading function, including wind percentile variable names
klawa_data_sources_paths = {
    'no_adapt': {'path': klawa_base_data_path, 'wind_var_name': MAP_LEVEL_TO_WIND_VAR['No Adaptation']},
    'medium_adapt': {'path': klawa_v99_data_path, 'wind_var_name': MAP_LEVEL_TO_WIND_VAR['Medium Adaptation']},
    'high_adapt': {'path': klawa_v995_data_path, 'wind_var_name': MAP_LEVEL_TO_WIND_VAR['High Adaptation']}
}

# Load and process data (cached)
pydeck_buildings_gdf, full_buildings_gdf, center_lat, center_lon, display_return_periods, building_centroids, all_return_periods = load_and_process_data(
    wind_data_path, klawa_data_sources_paths, VULNERABILITY_PARAMETERS, GEOMETRY_SIMPLIFICATION_TOLERANCE
)

if pydeck_buildings_gdf is None:
    st.error("Data processing failed or no buildings found. Please check logs for details.")
    st.stop()

# --- Streamlit Sidebar/Controls (Left Panel) ---
try:
    st.sidebar.image(LOGO_PATH_1, use_container_width=True)
except FileNotFoundError:
    st.sidebar.error("Logo 1 image file not found. Please check the path in the script.")
except Exception as e:
    st.sidebar.error(f"Error loading Logo 1: {e}")

st.sidebar.markdown("---")
st.sidebar.title("Select Adaptation Level for Map Display")
map_adaptation_level = st.sidebar.radio(
    " ",
    ('No Adaptation', 'Medium Adaptation', 'High Adaptation'),
    index=0
)

color_column_for_map = f'risk_color_rgba_{MAP_LEVEL_TO_SUFFIX[map_adaptation_level]}'

st.sidebar.markdown("---")
st.sidebar.header("Map Risk Zone Definition")
st.sidebar.write(f"**Red Zone:** Risk Index >= {RED_ZONE_THRESHOLD_PERC}")
st.sidebar.write(f"**Orange Zone:** Risk Index {ORANGE_ZONE_THRESHOLD_PERC} to < {RED_ZONE_THRESHOLD_PERC}")
st.sidebar.write(f"**Yellow Zone:** Risk Index {YELLOW_ZONE_THRESHOLD_PERC} to < {ORANGE_ZONE_THRESHOLD_PERC}")
st.sidebar.write(f"**Green Zone:** Risk Index < {YELLOW_ZONE_THRESHOLD_PERC}")

# --- Main Content Area Layout ---
input_col_1, input_col_2 = st.columns([0.5, 0.5])

with input_col_1:
    st.header("RiskX: Asset level risk analysis")
    st.subheader("Analyze a Single Address")
    address_query = st.text_input("Enter an address in Bergen:", key="address_input_top_col")

    if 'selected_building_id' not in st.session_state:
        st.session_state.selected_building_id = None
    if 'scenario_type' not in st.session_state:
        st.session_state.scenario_type = 'Moderate'
    if 'asset_value' not in st.session_state:
        st.session_state.asset_value = None

    st.session_state.asset_value = st.number_input(
        "Monetary Asset Value in million Norwegian Krone (NOK, Optional):",
        min_value=0.0,
        value=st.session_state.asset_value,
        format="%.2f",
        key="asset_value_input_top",
        help="Enter the monetary value of the asset (in millions of NOK) to calculate potential damages. This applies to the single address analysis."
    )
    if st.session_state.asset_value is not None and st.session_state.asset_value == 0.0:
        st.session_state.asset_value = None

with input_col_2:
    st.header(" ")
    st.subheader("Upload a File for Batch Analysis")
    uploaded_file = st.file_uploader("Upload an Excel file (columns: 'Address', 'Monetary Asset Value (million NOK)')", type=["xlsx"])

# --- Process Address Query (Single Address) and Display Odometers ---
if address_query:
    location = geocode_address(address_query, user_agent='bergen_wind_app_user_query')
    if location:
        query_point = (location.latitude, location.longitude)
        query_coords = np.array([[query_point[1], query_point[0]]]) # Longitude, Latitude for cKDTree
        distances, indices = cKDTree(np.array([[p.x, p.y] for p in building_centroids])).query(query_coords)

        nearest_building_index = indices[0]
        selected_building_gdf_row = full_buildings_gdf.iloc[nearest_building_index]
        st.session_state.selected_building_id = selected_building_gdf_row.name
    else:
        st.error("Could not geocode the address or find a nearby building. Please try a more specific address in Bergen.")
        st.session_state.selected_building_id = None

if st.session_state.selected_building_id is not None and not uploaded_file:
    selected_building_data = full_buildings_gdf.loc[st.session_state.selected_building_id]

    st.subheader("Risk Index & Wind Speed by Adaptation Level for Selected Building")
    
    klawa_levels_display_order = ['No Adaptation', 'Medium Adaptation', 'High Adaptation']

    odometer_cols = st.columns(3)

    for i, level_name in enumerate(klawa_levels_display_order):
        with odometer_cols[i]:
            klawa_index_norm = selected_building_data[f'normalized_klawa_risk_index_{MAP_LEVEL_TO_SUFFIX[level_name]}']
            wind_speed_value = selected_building_data[f'wind_speed_percentile_{MAP_LEVEL_TO_SUFFIX[level_name]}']

            st.markdown(f"**{level_name}**")
            if pd.isna(klawa_index_norm):
                st.markdown("<div style='text-align: center; color: grey;'>No Data</div>", unsafe_allow_html=True)
            else:
                gauge_color = get_gauge_color_hex(klawa_index_norm)

                unique_class_suffix = level_name.lower().replace(" ", "_").replace("(", "").replace(")", "").replace("-", "_")
                container_class = f"gauge-container-{unique_class_suffix}"
                fill_class = f"gauge-fill-{unique_class_suffix}"
                value_label_class = f"gauge-value-label-{unique_class_suffix}"

                st.markdown(
                    f"""
                    <style>
                    .{container_class}{{
                        width: 100%;
                        height: 25px;
                        background-color: #eee;
                        border-radius: 12px;
                        overflow: hidden;
                        position: relative;
                        margin-bottom: 5px;
                        box-shadow: inset 0 0 3px rgba(0,0,0,0.1);
                    }}
                    .{fill_class} {{
                        height: 100%;
                        border-radius: 12px;
                        background-color: {gauge_color};
                        width: {klawa_index_norm:.2f}%;
                        transition: width 0.5s ease-in-out;
                    }}
                    .{value_label_class} {{
                        position: absolute;
                        top: 50%;
                        left: 50%;
                        transform: translate(-50%, -50%);
                        font-weight: bold;
                        color: black;
                        text-shadow: 1px 1px 1px rgba(255,255,255,0.7);
                        font-size: 0.9em;
                    }}
                    </style>
                    <div class="{container_class}">
                        <div class="{fill_class}"></div>
                        <div class="{value_label_class}">{klawa_index_norm:.1f}</div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            
# --- Process Uploaded File (Batch Analysis) ---
if uploaded_file is not None:
    st.subheader("Batch Analysis Results")
    try:
        df_uploaded = pd.read_excel(uploaded_file)
        
        required_cols = ['Address', 'Monetary Asset Value (million NOK)']
        if not all(col in df_uploaded.columns for col in required_cols):
            st.error(f"Uploaded Excel must contain columns: {required_cols}")
        else:
            batch_results = []
            total_portfolio_asset_value = 0.0
            
            # Initializing Nominatim geolocator inside the try block
            geolocator = Nominatim(user_agent='bergen_wind_app_batch_query')

            with st.spinner("Processing addresses from Excel... This may take a while for many addresses."):
                for index, row in df_uploaded.iterrows():
                    address = row['Address']
                    asset_value_millions = row['Monetary Asset Value (million NOK)']
                    total_portfolio_asset_value += asset_value_millions * 1_000_000 # Accumulate total portfolio value

                    location = None
                    try:
                        location = geolocator.geocode(address, timeout=10)
                    except (GeocoderTimedOut, GeocoderServiceError) as e:
                        st.warning(f"Geocoding error for '{address}': {e}. Skipping this address.")

                    if location:
                        query_point = (location.latitude, location.longitude)
                        query_coords_for_batch = np.array([[query_point[1], query_point[0]]]) # Longitude, Latitude for cKDTree
                        
                        # Handle potential empty building_centroids if previous filters left no buildings
                        if building_centroids.empty:
                            st.error("No valid building data loaded for spatial matching. Please check input files and data ranges.")
                            break # Exit the loop if no buildings to match against

                        distances, indices = cKDTree(np.array([[p.x, p.y] for p in building_centroids])).query(query_coords_for_batch)
                        
                        nearest_building_index = indices[0]
                        
                        selected_building_data_batch = full_buildings_gdf.iloc[nearest_building_index]

                        # Calculate damage for 'Moderate' resilience, 100-year RP for summary
                        damage_ratio_moderate_rp100 = selected_building_data_batch.get(f'damage_ratio_moderate_rp{SELECTED_RP_FOR_DAMAGE_PLOT}', np.nan)
                        
                        calculated_damage_nok = np.nan
                        if pd.notna(damage_ratio_moderate_rp100) and pd.notna(asset_value_millions):
                            calculated_damage_nok = damage_ratio_moderate_rp100 * asset_value_millions * 1_000_000

                        # Calculate AAL for batch results (using 'Moderate' resilience for simplicity)
                        aal_nok = calculate_aal(selected_building_data_batch, 'Moderate', all_return_periods, asset_value_millions * 1_000_000)

                        batch_results.append({
                            'Address': address,
                            'Latitude': f"{location.latitude:.4f}",
                            'Longitude': f"{location.longitude:.4f}",
                            'Asset Value': f"{asset_value_millions * 1_000_000:,.2f}", # Store full NOK value here
                            'Risk Index (No Adapt)': f"{selected_building_data_batch['normalized_klawa_risk_index_no_adapt']:.2f}",
                            'Risk Index (Medium Adapt)': f"{selected_building_data_batch['normalized_klawa_risk_index_medium_adapt']:.2f}",
                            'Risk Index (High Adapt)': f"{selected_building_data_batch['normalized_klawa_risk_index_high_adapt']:.2f}",
                            f'Damage Ratio (RP {SELECTED_RP_FOR_DAMAGE_PLOT} yr, Moderate)': f"{(damage_ratio_moderate_rp100 * 100):.2f}%" if pd.notna(damage_ratio_moderate_rp100) else "N/A",
                            'Calculated Damage (NOK)': f"{calculated_damage_nok:,.2f}" if pd.notna(calculated_damage_nok) else "N/A",
                            'Average Annual Loss (NOK)': f"{aal_nok:,.2f}" if pd.notna(aal_nok) else "N/A",
                            'Map Risk Zone (No Adapt)': selected_building_data_batch['risk_zone_no_adapt'] # This will be 'Adaptation Level' in report
                        })
                    else:
                        batch_results.append({
                            'Address': address,
                            'Latitude': 'N/A', 'Longitude': 'N/A',
                            'Asset Value': 'N/A', # Add Asset Value for consistency
                            'Risk Index (No Adapt)': 'N/A',
                            'Risk Index (Medium Adapt)': 'N/A',
                            'Risk Index (High Adapt)': 'N/A',
                            f'Damage Ratio (RP {SELECTED_RP_FOR_DAMAGE_PLOT} yr, Moderate)': 'N/A',
                            'Calculated Damage (NOK)': 'N/A',
                            'Average Annual Loss (NOK)': 'N/A',
                            'Map Risk Zone (No Adapt)': 'N/A',
                            'Notes': 'Could not geocode address or find building'
                        })
            
            df_batch_results = pd.DataFrame(batch_results)
            st.dataframe(df_batch_results)

            # --- DOCX Report Download Button ---
            st.markdown("---")
            st.subheader("Generate Portfolio Report")
            
            # Use io.BytesIO to create a file-like object in memory
            docx_bytes = generate_docx_report(df_batch_results, len(df_uploaded), total_portfolio_asset_value)
            
            st.download_button(
                label="Download DOCX Report",
                data=docx_bytes,
                file_name="Portfolio_Wind_Risk_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                help="Generates a DOCX report of the batch analysis results."
            )

    except Exception as e:
        st.error(f"Error reading or processing Excel file: {e}")

st.markdown("---")

# --- Map and Damage Ratio Table Side-by-Side ---
map_col, damage_table_col = st.columns([0.6, 0.4])

with map_col:
    # --- Pydeck Map Setup ---
    view_state = pdk.ViewState(
        latitude=center_lat,
        longitude=center_lon,
        zoom=12,
        pitch=0,
    )

    layer = pdk.Layer(
        "GeoJsonLayer",
        pydeck_buildings_gdf,
        pickable=True,
        auto_highlight=True,
        get_fill_color=color_column_for_map,
        get_line_color=[0, 0, 0, 80],
        get_line_width=2,
        line_width_min_pixels=1,
    )

    r = pdk.Deck(
        map_style="mapbox://styles/mapbox/light-v9",
        initial_view_state=view_state,
        layers=[layer],
        tooltip={"text": "{tooltip}", "html": "{tooltip}"},
    )

    st.pydeck_chart(r, height=450)

    # --- Dynamic Legend ---
    st.subheader(f"Map Risk Zone Legend (Risk Index - {map_adaptation_level})")
    colors = ['#D4EDDA', '#FFD700', '#FB6A4A', '#DE2D26']
    labels = ['Low Risk', 'Medium-Low Risk', 'Medium-High Risk', 'High Risk']
    st.markdown(
        f"""
        <div style="display: flex; flex-wrap: wrap; gap: 8px; margin-top: 10px;">
            {"".join([f'<div style="background-color: {c}; width: 60px; height: 20px; border-radius: 4px; display: inline-block;"></div><span style="font-size: 0.9em; margin-right: 10px;">{l}</span>' for c, l in zip(colors, labels)])}
        </div>
        """,
        unsafe_allow_html=True
    )

with damage_table_col:
    if st.session_state.selected_building_id is not None and not uploaded_file: 
        st.subheader("Damage Ratio")

        st.write("Select Building Resilience:")
        resilience_cols_buttons = st.columns(3)
        
        with resilience_cols_buttons[0]:
            if st.button('Weak', key='res_weak', help="Assume weak building resilience"):
                st.session_state.scenario_type = 'Weak'
        with resilience_cols_buttons[1]:
            if st.button('Moderate', key='res_moderate', help="Assume moderate building resilience"):
                st.session_state.scenario_type = 'Moderate'
        with resilience_cols_buttons[2]:
            if st.button('Strong', key='res_strong', help="Assume strong building resilience"):
                st.session_state.scenario_type = 'Strong'

        current_scenario_type = st.session_state.scenario_type
        st.write(f"--- Results for **{current_scenario_type}** ---")

        scenario_results = []
        
        for rp in display_return_periods: 
            wind_speed = selected_building_data[f'wind_speed_rp{rp}']
            damage_ratio_col_name = f'damage_ratio_{current_scenario_type.lower()}_rp{rp}' 
            damage_ratio = selected_building_data[damage_ratio_col_name]
            
            row_data = {
                'Return Period (years)': rp,
                'Wind Speed (m/s)': f"{wind_speed:.2f}",
                'Damage Ratio (%)': f"{(damage_ratio * 100):.2f}%" if pd.notna(damage_ratio) else "N/A"
            }
            
            if st.session_state.asset_value is not None:
                calculated_damage = damage_ratio * st.session_state.asset_value * 1_000_000 if pd.notna(damage_ratio) else np.nan
                row_data['Calculated Damage (NOK)'] = f"{calculated_damage:,.2f}" if pd.notna(calculated_damage) else "N/A"
            
            scenario_results.append(row_data)

        st.markdown(pd.DataFrame(scenario_results).to_html(index=False), unsafe_allow_html=True)

        if st.session_state.asset_value is not None:
            aal_for_display = calculate_aal(
                selected_building_data, 
                current_scenario_type, 
                all_return_periods,
                st.session_state.asset_value * 1_000_000
            )
            st.markdown(f"---")
            if pd.notna(aal_for_display):
                st.metric(label=f"Average Annual Loss (AAL) for {current_scenario_type} Resilience", value=f"{aal_for_display:,.2f} NOK")
            else:
                st.info("AAL not calculated. Ensure asset value is provided and damage data is available.")

    elif uploaded_file is None:
        st.info("Enter an address to view damage ratio analysis.")


st.markdown("---")
st.caption("Powered by Streamlit, Pydeck, Xarray, OSMnx, Geopy, and Matplotlib.")
